from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "plugin-guides" / "CopiMineEndEvent_Command_Reference.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
BORDER = "B8C2CF"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total = sum(widths)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    ppr.append(num_pr)


def add_rich_paragraph(doc, parts, style=None, before=0, after=6, line=1.25):
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_spacing(paragraph, before, after, line)
    for text, options in parts:
        run = paragraph.add_run(text)
        set_run_font(run, **options)
    return paragraph


def add_list_paragraph(doc, text, num_id, bold_prefix=None):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=4, line=1.25)
    apply_numbering(paragraph, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_explicit_numbered_paragraph(doc, number, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=4, line=1.25)
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    run = paragraph.add_run(f"{number}. ")
    set_run_font(run, bold=True, color=INK)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_code_paragraph(doc, command, description):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=5, line=1.15)
    code = paragraph.add_run(command)
    set_run_font(code, name="Consolas", size=9.5, color=INK, bold=True)
    if description:
        detail = paragraph.add_run("  —  " + description)
        set_run_font(detail, size=10, color=MUTED)
    return paragraph


def add_command_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3672, 5688])
    set_table_borders(table)
    headers = ("Команда", "Назначение и безопасное поведение")
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.1)
        run = p.add_run(value)
        set_run_font(run, size=9.5, color=INK, bold=True)
    prevent_row_split(table.rows[0])
    repeat_table_header(table.rows[0])
    for command, description in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell)
        p = cells[0].paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.05)
        run = p.add_run(command)
        set_run_font(run, name="Consolas", size=9, color=INK, bold=True)
        p = cells[1].paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.1)
        run = p.add_run(description)
        set_run_font(run, size=9.5)
        prevent_row_split(table.rows[-1])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.05)
        run = p.add_run(value)
        set_run_font(run, size=9, color=INK, bold=True)
    prevent_row_split(table.rows[0])
    repeat_table_header(table.rows[0])
    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for cell, value in zip(row.cells, values):
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.05)
            run = p.add_run(str(value))
            set_run_font(run, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, heading, body, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D4DCE6", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=3, line=1.15)
    run = p.add_run(heading)
    set_run_font(run, size=10.5, color=INK, bold=True)
    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.15)
    run = p.add_run(body)
    set_run_font(run, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_keep_with_next(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return paragraph


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
    if "Guide Code" not in [style.name for style in doc.styles]:
        code = doc.styles.add_style("Guide Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        code.font.size = Pt(9.5)
        code.font.color.rgb = rgb(INK)
        code.paragraph_format.space_before = Pt(0)
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.line_spacing = 1.05


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=0, line=1.0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("COPIMINE  /  END RIFT EVENT")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = p.add_run("\tЛОКАЛЬНЫЙ СПРАВОЧНИК")
    set_run_font(right, size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(fp, after=0, line=1.0)
    run = fp.add_run("Только local runtime  •  страница ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field = OxmlElement("w:r")
    field.append(fld_begin)
    field.append(instr)
    field.append(fld_sep)
    field.append(fld_end)
    fp._p.append(field)


def add_title_block(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=3, line=1.0)
    run = p.add_run("END RIFT EVENT")
    set_run_font(run, size=24, color=INK, bold=True)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=15, line=1.0)
    run = p.add_run("Командный справочник для локальной проверки плагина")
    set_run_font(run, size=14, color=MUTED)
    metadata = (
        ("Среда", "environment: local; изолированный Paper + PostgreSQL"),
        ("Назначение", "настройка Core, диагностика волн, босса, рун, границы и gate"),
        ("Production", "не используется; launcher, сайт, боевой мир и данные игроков не затрагиваются"),
        ("Версия документа", "24 августа 2026"),
    )
    for label, value in metadata:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=2, line=1.05)
        lrun = p.add_run(label + ": ")
        set_run_font(lrun, size=10.5, color=INK, bold=True)
        vrun = p.add_run(value)
        set_run_font(vrun, size=10.5, color=INK)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=12, line=1.0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BLUE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def main():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title_block(doc)
    add_callout(
        doc,
        "Осторожно: команды с confirm меняют только локальный event state.",
        "Не запускай этот справочник на production. Core remove, resources reset, ritual reset, cleanup и gate restore/open рассчитаны на изолированный тестовый сервер. Ни одна из этих операций не должна выполняться на боевом сервере без отдельного плана, backup и ручного подтверждения.",
    )

    add_heading(doc, "1. Быстрый локальный прогон", 1)
    decimal = add_numbering(doc, "decimal")
    for text in (
        "Встань в локальном event world, наведи прицел на любой реальный твёрдый блок в радиусе 8 и выполни /cmend core set 2.",
        "Проверь /cmend status: ожидаются coreOverlay=true, coreModel=830001 и runes=2/2; ванильный block data исходного блока сохраняется.",
        "Собери ресурсы обычным взаимодействием с Core либо в тестовом состоянии используй /cmend resources add <MATERIAL> <amount>.",
        "Когда ресурсы заполнены, проверь READY_FOR_PLAYERS и overlays, затем войди на руны двумя survival-игроками.",
        "В Creative выполни /cmend test run creative: прогон волн, мини-боссов, всех boss phases, полёта spell-проектилей и cleanup не меняет official roster.",
        "Для точечной проверки gate задай две точки, выполни /cmend gate preview, затем /cmend gate open 1; после проверки верни блоки через /cmend gate restore confirm.",
    ):
        add_list_paragraph(doc, text, decimal)

    add_heading(doc, "2. Права и формат команд", 1)
    add_simple_table(
        doc,
        ("Permission node", "По умолчанию", "Что разрешает"),
        (
            ("copimine.endevent.admin", "op", "Core, arena, gate, ресурсы, ritual, manual wave/boss/client и cleanup."),
            ("copimine.endevent.test", "op", "Локальные disposable test wave/AI/boss/music и Creative full-run."),
            ("/cmend status", "без отдельного node", "Чтение текущего состояния event."),
        ),
        [2850, 1350, 5160],
    )
    add_callout(
        doc,
        "Подтверждение обязательно.",
        "Опасные операции намеренно не выполняются без слова confirm: /cmend core remove confirm, /cmend cleanup confirm, /cmend reset confirm, /cmend resources reset confirm, /cmend ritual reset confirm, /cmend unlock confirm и /cmend gate restore confirm.",
    )

    add_heading(doc, "3. Core, arena и визуальная граница", 1)
    add_command_table(
        doc,
        (
            ("/cmend core set <N>", "Игрок наводит прицел на твёрдый блок ≤8 блоков. Именно этот блок становится Core; его vanilla block data сохраняется, материал не заменяется."),
            ("/cmend core info", "То же состояние, что status: Core, ресурсы, руны, wave, boss и visual models."),
            ("/cmend core rebuild", "Пересобирает overlay Core и rune displays без сброса ресурсов."),
            ("/cmend core remove confirm", "Жёсткая граница сессии: восстанавливает исходный блок и pads, удаляет event-owned entities всех generations, projectiles, displays и client effects."),
            ("/cmend arena pos1|pos2", "Сохраняет пользовательские углы bounded арены в event world; блоки не меняются."),
            ("/cmend arena info", "Показывает координаты и объём текущей арены."),
            ("/cmend arena border [seconds]", "Показывает линиями частиц границу ±20 по X/Z и ±3 по Y. По умолчанию 10 секунд; мир не изменяется."),
            ("/cmend arena boundary <seconds>", "Алиас border."),
            ("/cmend arena clear confirm", "Сбрасывает custom arena и возвращает bounded bounds от Core."),
            ("/cmend portalroom set|info", "Сохраняет/показывает точку портальной комнаты; Core и арена не меняются."),
        ),
    )

    add_heading(doc, "4. Руны и ресурсы", 1)
    add_simple_table(
        doc,
        ("Ресурс", "Требование", "Цвет текста"),
        (
            ("DIAMOND / Алмазы", "100", "голубой / aqua"),
            ("ENDER_EYE / Око Эндера", "64", "зелёный"),
            ("BLAZE_ROD / Огненные стержни", "64", "золотой"),
            ("AMETHYST_SHARD / Осколки аметиста", "128", "светло-фиолетовый"),
        ),
        [2850, 1600, 4910],
    )
    add_command_table(
        doc,
        (
            ("/cmend resources status", "Показывает локальный прогресс по четырём ресурсам."),
            ("/cmend resources add <MATERIAL> <amount>", "Административный локальный helper; amount ограничивается требованием и не выдаёт предметы игроку."),
            ("/cmend resources reset confirm", "Обнуляет прогресс вне боя и возвращает фазу в COLLECTING."),
            ("Обычное взаимодействие с Core", "Боевой путь сдачи ресурсов игроком. После полного заполнения появляются/пересобираются руны и Core переходит в READY_FOR_PLAYERS."),
        ),
    )
    add_callout(
        doc,
        "Занятая руна меняет display model.",
        "В status поле visuals показывает runes=2/2 occupied=0 или occupied=1/2. Survival-игрок на pad переводит overlay в occupied-вариант; выход игрока возвращает свободный вариант. Pads остаются на полу над реальными vanilla-блоками, а не летают в воздухе.",
    )

    add_heading(doc, "5. Gate: particle preview и послойное открытие", 1)
    add_command_table(
        doc,
        (
            ("/cmend gate pos1", "Сохраняет первую точку и на 10 секунд рисует particle-only highlight. Блоки не заменяются."),
            ("/cmend gate pos2", "Сохраняет вторую точку, проверяет один world и bounded volume, снова рисует preview."),
            ("/cmend gate info", "Показывает pos1, pos2, PREVIEW/OPENING/OPENED/RESTORED и progress."),
            ("/cmend gate preview", "Сохраняет durable snapshot выбранного cuboid и оставляет все block materials нетронутыми."),
            ("/cmend gate open [ticks-per-layer]", "Удаляет только точный inclusive cuboid сверху вниз по слоям с проверкой snapshot; conflicts не перезаписываются."),
            ("/cmend gate restore confirm", "Возвращает snapshot после локального теста gate; selection points сохраняются для повторной проверки."),
        ),
    )
    add_callout(
        doc,
        "Размер gate не связан с границей арены.",
        "Арена всегда bounded вокруг Core: ±20 по горизонтали, ±3 по вертикали. Gate — отдельный выбранный cuboid. Для victory opening используется тот же durable snapshot и тот же top-down animation; если snapshot конфликтует с чужим блоком, слой не перезаписывается и операция прерывается.",
    )

    add_heading(doc, "6. Ритуал и безопасное восстановление", 1)
    add_command_table(
        doc,
        (
            ("/cmend ritual start", "Запускает проверку/сбор игроков только в READY_FOR_PLAYERS."),
            ("/cmend ritual cancel confirm", "Отменяет countdown; текущий бой не используется как способ cleanup."),
            ("/cmend ritual cleanup confirm", "Очищает transient event-owned entities и client effects текущей сессии."),
            ("/cmend ritual reset confirm", "Сбрасывает event state, но не трогает End world, player data и DB."),
            ("/cmend ritual unlock confirm", "Административная проверка unlock; не используй на production."),
            ("/cmend cleanup confirm", "Очищает event-owned entities текущей сессии без общего world scan."),
            ("/cmend reset confirm", "Короткий alias безопасного event reset с защитой от UNLOCKED."),
            ("/cmend unlock confirm", "Административный unlock helper; обычный victory saga выполняет unlock сам."),
        ),
    )

    add_heading(doc, "7. Волны, ИИ и боссы", 1)
    add_command_table(
        doc,
        (
            ("/cmend wave spawn <1|2|3|final>", "Создаёт disposable wave controller; official victory roster не меняется."),
            ("/cmend wave clear", "Удаляет только event-owned wave entities."),
            ("/cmend test wave <1|2|3|final>", "То же через test permission и с явной пометкой локального теста."),
            ("/cmend test ai", "Запускает реальный leash/target/teleport AI path без official session."),
            ("/cmend test boss", "Создаёт disposable boss path."),
            ("/cmend boss spawn", "Test boss; для official всегда требуется explicit confirm."),
            ("/cmend boss spawn official confirm", "Запускает official boss только с подтверждением; не используй в smoke-тесте."),
            ("/cmend boss info", "Показывает boss и фазу."),
            ("/cmend boss damage <n>", "Наносит test damage живому boss."),
            ("/cmend boss phase <normal|half|final>", "Переводит test boss между фазами."),
            ("/cmend boss kill cleanup", "Удаляет test boss без victory."),
            ("/cmend boss spell <void_blast|rift_projectile|void_mark|summon|control_reverse>", "Проверяет конкретный boss spell path; projectiles летят с particle telegraph/flight/cast."),
        ),
    )
    add_simple_table(
        doc,
        ("Компонент", "Настройка локального события"),
        (
            ("Волна 1", "5 endermen + 8 spiders"),
            ("Волна 2", "7 endermen + 6 spiders + 3 shulkers"),
            ("Волна 3", "10 endermen + 3 shulkers + 4 elite endermen"),
            ("Final wave", "8 spiders + 2 shulkers + 6 elite endermen"),
            ("Endermites", "Не спавнятся; заменены spiders."),
            ("Spider tuning", "16 base HP + 10 = 26 HP; 2 base damage + 2 = 4 damage."),
            ("Arena AI", "Horizontal range ≤20; teleport destination keeps Core Y; above/below Core is forbidden. Core block/top position is allowed to prevent crowding."),
            ("Boss", "1200 HP; half phase 600; final threshold 120; final ritual health 200."),
        ),
        [2500, 6860],
    )

    add_heading(doc, "8. Названия заклинаний", 1)
    add_simple_table(
        doc,
        ("ID", "Русское название", "Кто использует"),
        (
            ("void_blast", "Взрыв Бездны", "Boss"),
            ("rift_projectile", "Снаряд Разлома", "Boss"),
            ("void_mark", "Клеймо Пустоты", "Boss"),
            ("summon_servants", "Призыв слуг Разлома", "Boss"),
            ("will_distortion", "Искажение воли", "Boss, стадия ниже 50%"),
            ("rift_step", "Рывок Разлома", "Mini-boss elite"),
            ("void_snare", "Кандалы Пустоты", "Mini-boss elite"),
            ("echo_pulse", "Импульс Эха", "Mini-boss elite"),
        ),
        [2400, 4100, 2860],
    )
    add_callout(
        doc,
        "Проверяй не только текст telegraph.",
        "В Paper latest.log для одного полного run должны появиться BOSS_SPELL_TELEGRAPH → BOSS_SPELL_FLIGHT → BOSS_SPELL_CAST и MINIBOSS_SPELL_TELEGRAPH → MINIBOSS_SPELL_FLIGHT → MINIBOSS_SPELL_CAST. Creative runner ждёт flight/cast мини-боссов перед очисткой wave 3.",
    )

    add_heading(doc, "9. Музыка без слов", 1)
    add_simple_table(
        doc,
        ("Фаза", "Источник/название", "Sound id и длина"),
        (
            ("Волны", "cynicmusic — Battle Theme A (CC0)", "copimine:end_rift/waves · 95.85 s"),
            ("Boss 100–50%", "SubspaceAudio / Juhani Junkala — Boss Battle Music (CC0)", "copimine:end_rift/boss · 123.43 s"),
            ("Boss ниже 50%", "nene — Boss Battle 2: Symphonic Metal (CC0, instrumental)", "copimine:end_rift/boss_half · 26.48 s"),
            ("Boss ниже 10%", "cynicmusic — Dramatic Boss Encounter (CC0)", "copimine:end_rift/boss_final · 39.92 s"),
            ("Победа", "cynicmusic — Victory Theme for RPG (CC0)", "copimine:end_rift/victory · 20.00 s"),
        ),
        [1700, 4600, 3060],
    )
    add_paragraph = doc.add_paragraph()
    set_paragraph_spacing(add_paragraph, after=6, line=1.15)
    run = add_paragraph.add_run("Лицензии и исходные страницы: ")
    set_run_font(run, size=9.5, color=MUTED, bold=True)
    run = add_paragraph.add_run("resourcepacks/END_RIFT_MUSIC_LICENSES.md. Все пять треков в pack — инструментальные производные OGG; в resource pack нет lyric/vocal дорожек.")
    set_run_font(run, size=9.5, color=MUTED)

    add_heading(doc, "10. Client bridge и музыкальный smoke", 1)
    add_command_table(
        doc,
        (
            ("/cmend client status [player]", "Показывает channel, boss-id, control-id и активные client bindings."),
            ("/cmend client bindboss [player]", "Повторно отправляет boss/control binding выбранному игроку или всем online."),
            ("/cmend client clear [player]", "Очищает только client effects event."),
            ("/cmend test music <waves|boss|half|final|victory> <player>", "Включает выбранный трек на online player и пишет END_EVENT_MUSIC_TEST в лог."),
            ("/cmend debug или /cmend recovery", "Read-only alias status для диагностики."),
        ),
    )

    add_heading(doc, "11. Минимальный набор доказательств перед фиксацией", 1)
    evidence_items = (
        "python -m pytest -q tests — contracts должны быть зелёными; отдельный RED-тест должен существовать для каждого исправляемого runtime-дефекта.",
        ".\\tests\\RunEndRiftEventChecks.ps1 — builds WorldCore/Artifacts/End Event/Client, resource pack и Python/Java contracts.",
        ".\\tests\\StartEndRiftLocal.ps1 — только local-runtime/end-rift-server на 25566/25576 с local PostgreSQL на 55433.",
        "Проверить latest.log: services ready, core visuals, resource completion, spider stats, boss 1200 HP, spell flight/cast, creative cleanup и отсутствие NoClassDefFoundError.",
        "После destructive-local теста выполнить /cmend core remove confirm и убедиться в status: UNCONFIGURED, event-mobs=0, boss=none, coreOverlay=false, runes=0/0.",
        "Проверить git status, commit и push только в codex/end-rift-event; production deploy и launcher/site изменения в этой ветке не выполняются.",
    )
    for number, text in enumerate(evidence_items, start=1):
        add_explicit_numbered_paragraph(doc, number, text)

    add_heading(doc, "12. Полезные runtime markers", 1)
    add_simple_table(
        doc,
        ("Marker", "Что доказывает"),
        (
            ("END_EVENT_PHYSICAL_VISUALS", "Core overlay и rune overlays созданы поверх vanilla floor/block data."),
            ("END_EVENT_GATE_SELECTION_PREVIEW", "Выбранный gate подсвечен частицами и generation-bound task запущен."),
            ("END_EVENT_GATE_LAYER", "Очередной точный слой gate обработан сверху вниз."),
            ("END_EVENT_OWNED_CLEANUP", "Удалены event-owned entities across generations."),
            ("SPIDER_STATS", "Фактические spider health/attack после tuning."),
            ("BOSS_SPELL_FLIGHT / MINIBOSS_SPELL_FLIGHT", "Spell реально прошёл flight phase, а не только показал текст."),
            ("CREATIVE_TEST_COMPLETE", "Full-run завершён и official phase/roster не изменены."),
        ),
        [3600, 5760],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
