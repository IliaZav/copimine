from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "plugin-guides"


PLUGIN_SPECS = [
    {
        "id": "admin-plus",
        "name": "CopiMineUltimateAdminPlus",
        "owner": "Админ-хаб и делегатор",
        "plugin_yml": ROOT / "copimine-admin-plugin" / "plugin.yml",
        "java": ROOT / "copimine-admin-plugin" / "src" / "me" / "copimine" / "ultimateplus" / "CopiMineUltimateAdminPlus.java",
    },
    {
        "id": "economy-core",
        "name": "CopiMineEconomyCore",
        "owner": "Экономика, банк, ATM, PIN и donation foundation",
        "plugin_yml": ROOT / "copimine-economy-core" / "plugin.yml",
        "java": ROOT / "copimine-economy-core" / "src" / "me" / "copimine" / "economycore" / "CopiMineEconomyCore.java",
    },
    {
        "id": "election-core",
        "name": "CopiMineElectionCore",
        "owner": "Выборы, участки, ЦИК, бюллетени, президент и live flow",
        "plugin_yml": ROOT / "copimine-election-core" / "plugin.yml",
        "java": ROOT / "copimine-election-core" / "src" / "me" / "copimine" / "electioncore" / "CopiMineElectionCore.java",
    },
    {
        "id": "artifacts",
        "name": "CopiMineArtifacts",
        "owner": "AR-лавка, donation item issuance, repair, reclaim, anti-dupe",
        "plugin_yml": ROOT / "copimine-artifacts" / "plugin.yml",
        "java": ROOT / "copimine-artifacts" / "src" / "me" / "copimine" / "artifacts" / "CopiMineArtifacts.java",
    },
    {
        "id": "narcotics",
        "name": "CopiMineNarcotics",
        "owner": "Наркотики, варка в котле, overdose и visuals/client bridge",
        "plugin_yml": ROOT / "copimine-narcotics" / "plugin.yml",
        "java": ROOT / "copimine-narcotics" / "src" / "me" / "copimine" / "narcotics" / "CopiMineNarcotics.java",
    },
    {
        "id": "world-core",
        "name": "CopiMineWorldCore",
        "owner": "Миры, border, доступ к Nether/End и world control",
        "plugin_yml": ROOT / "copimine-world-core" / "plugin.yml",
        "java": ROOT / "copimine-world-core" / "src" / "me" / "copimine" / "worldcore" / "CopiMineWorldCore.java",
    },
]


@dataclass
class ButtonInfo:
    label: str
    action: str


@dataclass
class ScreenInfo:
    method: str
    title: str
    buttons: list[ButtonInfo] = field(default_factory=list)


def set_run_font(run, name: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document, text: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, "Calibri", 22, True)
    if subtitle:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(10)
        sub_run = sub.add_run(subtitle)
        set_run_font(sub_run, "Calibri", 11, False)


def load_plugin_meta(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    result: dict[str, object] = {"commands": {}}
    current_section: str | None = None
    current_command: str | None = None
    current_permission: str | None = None

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line or line.lstrip().startswith("#"):
            continue
        indent = len(line) - len(line.lstrip(" "))
        stripped = line.strip()

        if indent == 0 and ":" in stripped:
            key, value = stripped.split(":", 1)
            key = key.strip()
            value = value.strip().strip("'\"")
            current_section = key
            current_command = None
            current_permission = None
            if key in {"depend", "softdepend", "libraries"}:
                result[key] = []
            elif key not in {"commands", "permissions"}:
                result[key] = value
            continue

        if current_section == "commands":
            if indent == 2 and stripped.endswith(":"):
                current_command = stripped[:-1].strip()
                result["commands"][current_command] = {}
                continue
            if indent >= 4 and current_command and ":" in stripped:
                key, value = stripped.split(":", 1)
                key = key.strip()
                value = value.strip()
                if key == "aliases":
                    value = value.strip("[]")
                    aliases = [item.strip().strip("'\"") for item in value.split(",") if item.strip()]
                    result["commands"][current_command][key] = aliases
                else:
                    result["commands"][current_command][key] = value.strip("'\"")
            continue

        if current_section == "permissions":
            if indent == 2 and stripped.endswith(":"):
                current_permission = stripped[:-1].strip()
                result.setdefault("permissions", {})[current_permission] = {}
                continue
            if indent >= 4 and current_permission and ":" in stripped:
                key, value = stripped.split(":", 1)
                result["permissions"][current_permission][key.strip()] = value.strip().strip("'\"")
            continue

        if current_section in {"depend", "softdepend", "libraries"} and stripped.startswith("-"):
            result.setdefault(current_section, []).append(stripped.lstrip("-").strip())

    return result


def clean_text(text: str) -> str:
    return re.sub(r"&[0-9a-fk-or]", "", text, flags=re.I).replace("&l", "").replace("&r", "").strip()


def find_matching_brace(text: str, start_index: int) -> int:
    depth = 0
    for index in range(start_index, len(text)):
        char = text[index]
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return index
    return -1


def extract_screens(java_path: Path) -> list[ScreenInfo]:
    text = java_path.read_text(encoding="utf-8")
    screens: list[ScreenInfo] = []
    fn_pattern = re.compile(
        r"(?:private|public)\s+void\s+"
        r"(open[A-Z][A-Za-z0-9_]*|legacyOpen[A-Z][A-Za-z0-9_]*|show[A-Z][A-Za-z0-9_]*)"
        r"\s*\([^)]*\)\s*(?:throws\s+[^{]+)?\{",
        re.S,
    )
    create_pattern = re.compile(r'create\(m,\s*\d+,\s*"([^"]+)"\)')
    btn_pattern = re.compile(
        r'btn\(m,\s*\d+,\s*Material\.[A-Z0-9_]+,\s*"([^"]*)",.*?,"([^"]*)"\)',
        re.S,
    )
    for match in fn_pattern.finditer(text):
        method = match.group(1)
        body_start = match.end() - 1
        body_end = find_matching_brace(text, body_start)
        if body_end < 0:
            continue
        body = text[body_start:body_end + 1]
        title_match = create_pattern.search(body)
        if not title_match:
            continue
        title = clean_text(title_match.group(1))
        buttons: list[ButtonInfo] = []
        for btn_match in btn_pattern.finditer(body):
            label = clean_text(btn_match.group(1))
            action = clean_text(btn_match.group(2))
            if label:
                buttons.append(ButtonInfo(label=label, action=action or "none"))
        screens.append(ScreenInfo(method=method, title=title, buttons=buttons))
    return screens


def add_command_table(doc: Document, commands: dict[str, dict]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Команда", "Описание", "Использование", "Aliases"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for command_name, payload in commands.items():
        row = table.add_row().cells
        row[0].text = f"/{command_name}"
        row[1].text = str(payload.get("description", "") or "")
        row[2].text = str(payload.get("usage", "") or "")
        aliases = payload.get("aliases") or []
        row[3].text = ", ".join(str(alias) for alias in aliases)


def add_screen_section(doc: Document, screen: ScreenInfo) -> None:
    doc.add_heading(screen.title or screen.method, level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("Метод: ").bold = True
    p.add_run(screen.method)
    if not screen.buttons:
        doc.add_paragraph("Кнопки не найдены автоматически.")
        return
    for button in screen.buttons:
        para = doc.add_paragraph(style="List Bullet")
        label_run = para.add_run(f"{button.label}")
        label_run.bold = True
        para.add_run(" — ")
        para.add_run(button.action or "none")


def chunked(iterable: Iterable[ScreenInfo], size: int) -> Iterable[list[ScreenInfo]]:
    bucket: list[ScreenInfo] = []
    for item in iterable:
        bucket.append(item)
        if len(bucket) == size:
            yield bucket
            bucket = []
    if bucket:
        yield bucket


def build_plugin_doc(spec: dict) -> Path:
    plugin_meta = load_plugin_meta(spec["plugin_yml"])
    screens = extract_screens(spec["java"])
    doc = Document()
    style_document(doc)
    add_title(
        doc,
        f"{spec['name']} — команды и GUI",
        f"Назначение: {spec['owner']}",
    )

    intro = doc.add_paragraph()
    intro.add_run("Главный класс: ").bold = True
    intro.add_run(str(plugin_meta.get("main", "")))
    intro.add_run(" | Версия: ").bold = True
    intro.add_run(str(plugin_meta.get("version", "")))

    if plugin_meta.get("depend") or plugin_meta.get("softdepend"):
        dep = doc.add_paragraph()
        dep.add_run("Зависимости: ").bold = True
        depends = list(plugin_meta.get("depend") or [])
        softdepends = list(plugin_meta.get("softdepend") or [])
        dep.add_run(", ".join(depends) if depends else "нет жёстких")
        if softdepends:
            dep.add_run(" | Softdepend: ").bold = True
            dep.add_run(", ".join(softdepends))

    doc.add_heading("Команды", level=1)
    commands = plugin_meta.get("commands") or {}
    if commands:
        add_command_table(doc, commands)
    else:
        doc.add_paragraph("В plugin.yml не заявлены команды.")

    doc.add_heading("GUI", level=1)
    if not screens:
        doc.add_paragraph("Автоматически найденных GUI-экранов не обнаружено.")
    else:
        summary = doc.add_paragraph()
        summary.add_run("Найдено экранов: ").bold = True
        summary.add_run(str(len(screens)))
        for pack in chunked(screens, 12):
            for screen in pack:
                add_screen_section(doc, screen)

    out_path = OUT_DIR / f"{spec['name']}.docx"
    doc.save(out_path)
    return out_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generated = [build_plugin_doc(spec) for spec in PLUGIN_SPECS]
    print("\n".join(str(path) for path in generated))


if __name__ == "__main__":
    main()
