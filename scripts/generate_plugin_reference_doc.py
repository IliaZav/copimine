from __future__ import annotations

import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "plugin-guides"
OUTPUT_PATH = OUT_DIR / "CopiMine-plugins-command-reference.docx"
SERVER_PLUGINS_DIR = ROOT / "minecraft" / "server" / "plugins"


@dataclass
class CommandRow:
    command: str
    description: str
    usage: str = ""
    aliases: list[str] = field(default_factory=list)


@dataclass
class PluginRecord:
    name: str
    description: str
    version: str = ""
    main_class: str = ""
    jar_name: str = ""
    depends: list[str] = field(default_factory=list)
    softdepends: list[str] = field(default_factory=list)
    commands: list[CommandRow] = field(default_factory=list)


CUSTOM_PLUGIN_SPECS = [
    {
        "name": "CopiMineUltimateAdminPlus",
        "description": "Главный игровой админ-хаб. Открывает панели управления, даёт доступ к отчётам, проверкам игроков и делегирует экономику и выборы в профильные плагины.",
        "plugin_yml": ROOT / "copimine-admin-plugin" / "plugin.yml",
    },
    {
        "name": "CopiMineEconomyCore",
        "description": "Ядро экономики. Ведёт банковские счета, ATM, PIN, донатный баланс, платёжные сессии и денежные операции.",
        "plugin_yml": ROOT / "copimine-economy-core" / "plugin.yml",
    },
    {
        "name": "CopiMineElectionCore",
        "description": "Ядро выборов. Управляет стадиями выборов, участками, ЦИК, бюллетенями, президентским мандатом и сопутствующими панелями.",
        "plugin_yml": ROOT / "copimine-election-core" / "plugin.yml",
    },
    {
        "name": "CopiMineArtifacts",
        "description": "AR- и donation-лавки. Отвечает за каталог предметов, выдачу, ремонт, возврат утерянных вещей и защиту от дюпа.",
        "plugin_yml": ROOT / "copimine-artifacts" / "plugin.yml",
    },
    {
        "name": "CopiMineNarcotics",
        "description": "Система наркотиков, варки в котле, овердозов и визуальных эффектов с серверным fallback и CopiMineClient bridge.",
        "plugin_yml": ROOT / "copimine-narcotics" / "plugin.yml",
    },
    {
        "name": "CopiMineWorldCore",
        "description": "Управление мирами, границами, доступом в Nether и End, безопасными телепортами и статусом открытия миров.",
        "plugin_yml": ROOT / "copimine-world-core" / "plugin.yml",
    },
]


CUSTOM_COMMANDS = {
    "CopiMineUltimateAdminPlus": [
        CommandRow("/cmultra", "Открывает основной админ-хаб.", "/cmultra", ["cmplus", "uadmin", "ultraadmin", "adminka"]),
        CommandRow("/cmultra menu", "Открывает основной админ-хаб.", "/cmultra menu"),
        CommandRow("/cmultra ar sync", "Показывает сообщение о переносе AR-синхронизации в EconomyCore и открывает экономический хаб, если есть доступ.", "/cmultra ar sync"),
        CommandRow("/cmultra check start <player>", "Запускает админскую проверку игрока.", "/cmultra check start Nick"),
        CommandRow("/cmultra check stop <player>", "Завершает активную проверку игрока.", "/cmultra check stop Nick"),
        CommandRow("/cmultra check return <player>", "Возвращает игрока после проверки.", "/cmultra check return Nick"),
        CommandRow("/cadm", "Открывает короткий админ-хаб CopiMine.", "/cadm", ["copimine", "cadmin", "cpanel", "adminhub", "cpa", "cpadmin"]),
        CommandRow("/ar", "Открывает раздел экономики в админ-хабе, если доступен EconomyCore.", "/ar", ["ars", "arreconomy", "aradmin"]),
        CommandRow("/cmbank", "Открывает банковый раздел админ-хаба.", "/cmbank", ["bank", "cbank", "bankar"]),
        CommandRow("/report <текст>", "Отправляет жалобу или сообщение администрации.", "/report Текст сообщения", ["problem", "helpme", "ticket"]),
        CommandRow("/appeal <текст>", "Отправляет апелляцию или обращение в администрацию.", "/appeal Текст сообщения", ["adminrequest", "request", "helpadm"]),
        CommandRow("/rpguard status", "Показывает статус защиты RP-команд.", "/rpguard status", ["cmrpguard", "rpcommands"]),
        CommandRow("/rpguard test", "Проверяет текущую конфигурацию защиты RP-команд.", "/rpguard test", ["cmrpguard", "rpcommands"]),
    ],
    "CopiMineEconomyCore": [
        CommandRow("Командный вход через GUI", "У EconomyCore нет собственного публичного набора чат-команд. Игроки и админы работают через ATM, банковые меню и панели, которые открывает AdminPlus или игровые блоки."),
    ],
    "CopiMineElectionCore": [
        CommandRow("/hidelive", "Локально скрывает live-панель выборов у игрока.", "/hidelive"),
    ],
    "CopiMineArtifacts": [
        CommandRow("/cmartifacts", "Открывает админское меню артефактов или показывает игроку справку по лавке.", "/cmartifacts", ["cartifacts", "artifactshop"]),
        CommandRow("/cmartifacts claim", "Выдаёт игроку купленные или ожидающие предметы, если они готовы к получению.", "/cmartifacts claim"),
        CommandRow("/cmartifacts repair", "Открывает механику ремонта AR-предметов.", "/cmartifacts repair"),
        CommandRow("/cmartifacts shop", "Показывает краткую справку по управлению лавками.", "/cmartifacts shop"),
        CommandRow("/cmartifacts shop create <shop_id>", "Создаёт AR- или donation-лавку на выбранном блоке.", "/cmartifacts shop create ar_main"),
        CommandRow("/cmartifacts shop remove", "Удаляет лавку с выбранного блока.", "/cmartifacts shop remove"),
        CommandRow("/cmartifacts shop list", "Показывает список зарегистрированных лавок.", "/cmartifacts shop list"),
        CommandRow("/cmartifacts shop open <shop_id>", "Открывает витрину конкретной лавки.", "/cmartifacts shop open ar_main"),
        CommandRow("/cmartifacts admin", "Открывает админское меню артефактов.", "/cmartifacts admin"),
        CommandRow("/cmartifacts reload", "Перечитывает конфигурацию и каталоги артефактов.", "/cmartifacts reload"),
    ],
    "CopiMineNarcotics": [
        CommandRow("/cmnarcotics give <player> <id|all>", "Выдаёт официальный наркотик игроку.", "/cmnarcotics give Nick feta"),
        CommandRow("/cmnarcotics info <player>", "Показывает состояние игрока по весу и овердозу.", "/cmnarcotics info Nick"),
        CommandRow("/cmnarcotics setweight <id> <value>", "Меняет вес конкретного наркотика.", "/cmnarcotics setweight feta 120"),
        CommandRow("/cmnarcotics setthreshold <value>", "Меняет общий порог овердоза.", "/cmnarcotics setthreshold 500"),
        CommandRow("/cmnarcotics setwindow <seconds>", "Меняет окно накопления веса.", "/cmnarcotics setwindow 900"),
        CommandRow("/cmnarcotics setduration <seconds>", "Меняет длительность активного эффекта.", "/cmnarcotics setduration 120"),
        CommandRow("/cmnarcotics reload", "Перечитывает конфигурацию наркотиков.", "/cmnarcotics reload"),
        CommandRow("/cmnarcotics selfcheck", "Проверяет конфиг, runtime и ресурс-пак на ошибки.", "/cmnarcotics selfcheck"),
        CommandRow("/cmnarcotics reset confirm", "Асинхронно очищает narcotics-state и связанные таблицы плагина.", "/cmnarcotics reset confirm"),
        CommandRow("/cmnarcotics clearoverdose <player>", "Сбрасывает состояние овердоза у игрока.", "/cmnarcotics clearoverdose Nick"),
        CommandRow("/cmnarcotics texture mode vanilla", "Переключает текстурный режим на обычные vanilla-предметы.", "/cmnarcotics texture mode vanilla"),
        CommandRow("/cmnarcotics texture mode custom", "Включает режим custom-моделей для официальных предметов.", "/cmnarcotics texture mode custom"),
        CommandRow("/cmnarcotics texture migrate online", "Обновляет текстурный режим у официальных предметов у игроков онлайн.", "/cmnarcotics texture migrate online"),
        CommandRow("/cmnarcotics texture migrate nearby", "Обновляет официальные предметы в ближайших инвентарях и контейнерах.", "/cmnarcotics texture migrate nearby"),
        CommandRow("/cmnarcotics visuals status", "Показывает статус визуальных маршрутов: client mod, server overlay и fallback.", "/cmnarcotics visuals status"),
        CommandRow("/cmnarcotics visuals enable <effectId|all>", "Включает визуальный эффект или сразу весь набор.", "/cmnarcotics visuals enable DESATURATE"),
        CommandRow("/cmnarcotics visuals disable <effectId|all>", "Отключает визуальный эффект или сразу весь набор.", "/cmnarcotics visuals disable all"),
        CommandRow("/cmnarcotics visuals mode auto", "Возвращает автоматический выбор маршрута визуалов.", "/cmnarcotics visuals mode auto"),
        CommandRow("/cmnarcotics visuals mode client_mod", "Принудительно предпочитает CopiMineClient visual route.", "/cmnarcotics visuals mode client_mod"),
        CommandRow("/cmnarcotics visuals mode server_overlay", "Принудительно использует серверный overlay-маршрут.", "/cmnarcotics visuals mode server_overlay"),
        CommandRow("/cmnarcotics visuals mode server_fallback", "Принудительно использует только частицы и зелья.", "/cmnarcotics visuals mode server_fallback"),
        CommandRow("/cmnarcotics visuals test <player> <effectId> [seconds]", "Тестирует визуальный эффект на выбранном игроке.", "/cmnarcotics visuals test Nick CHAOS 10"),
        CommandRow("/cmclient check <player>", "Показывает, подключён ли CopiMineClient и какие возможности он отдал серверу.", "/cmclient check Nick"),
        CommandRow("/cmclient visualtest <player> <effectId> [seconds]", "Отправляет игроку client-side visual test через CopiMineClient.", "/cmclient visualtest Nick INVERT 8"),
        CommandRow("/cmclient fallbacktest <player> <effectId> [seconds]", "Запускает тот же эффект через серверный fallback для сравнения.", "/cmclient fallbacktest Nick INVERT 8"),
        CommandRow("/cmclient require client true|false", "Включает или выключает принудительное требование CopiMineClient.", "/cmclient require client false"),
    ],
    "CopiMineWorldCore": [
        CommandRow("/cmworld status", "Показывает общий статус миров, границы и доступность Nether/End.", "/cmworld status"),
        CommandRow("/cmworld reload", "Перечитывает конфигурацию WorldCore.", "/cmworld reload"),
        CommandRow("/cmworld safecheck", "Проверяет точки безопасного возврата и защитные fallback-правила.", "/cmworld safecheck"),
        CommandRow("/cmworld border status", "Показывает текущие ограничения border.", "/cmworld border status"),
        CommandRow("/cmworld border apply", "Применяет текущую конфигурацию border к мирам.", "/cmworld border apply"),
        CommandRow("/cmworld border set <radius> [confirm]", "Меняет радиус world border.", "/cmworld border set 5000 confirm"),
        CommandRow("/cmworld nether status", "Показывает состояние Nether.", "/cmworld nether status"),
        CommandRow("/cmworld nether open", "Открывает Nether.", "/cmworld nether open"),
        CommandRow("/cmworld nether close [confirm]", "Закрывает Nether после подтверждения.", "/cmworld nether close confirm"),
        CommandRow("/cmworld end status", "Показывает состояние End.", "/cmworld end status"),
        CommandRow("/cmworld end open", "Открывает End.", "/cmworld end open"),
        CommandRow("/cmworld end close [confirm]", "Закрывает End после подтверждения.", "/cmworld end close confirm"),
    ],
}


THIRD_PARTY_DESCRIPTIONS = {
    "AuthEffects": "Служебный плагин эффектов вокруг авторизации и входа игрока.",
    "AuthMe": "Авторизация и регистрация игроков на сервере.",
    "Chunky": "Предварительная генерация мира и контроль очередей предгена.",
    "CoreProtect": "Логи блоков, откаты и поиск действий по миру.",
    "emotecraft": "Эмоции и анимации игрока.",
    "EntityClearer": "Очистка лишних сущностей по таймеру и правилам.",
    "Essentials": "Базовые сервисные команды сервера: дом, телепорт, инвентарь, сообщения и утилиты.",
    "EssentialsChat": "Форматирование чата и связанные чат-команды Essentials.",
    "EssentialsSpawn": "Управление точкой спавна сервера и командами возрождения.",
    "FarmControl": "Ограничение лагов от ферм, мобов и сущностей.",
    "GrimAC": "Античит и служебные команды проверки.",
    "GSit": "Команды, чтобы сесть, лечь, ползать или позировать.",
    "ImageFrame": "Отображение картинок и карт на картах/рамках.",
    "LuckPerms": "Группы, роли, права и наследование разрешений.",
    "PlaceholderAPI": "Плейсхолдеры для чата, TAB, GUI и других интеграций.",
    "ProtocolLib": "Сетевой мост для других плагинов. Пользовательских команд обычно почти нет.",
    "SeeMore": "Служебный плагин улучшенного обзора и клиентских подсказок.",
    "TAB": "TAB-лист, префиксы, nametag и scoreboard-данные.",
    "Vault": "Интеграционный мост между экономикой, правами и другими плагинами.",
    "voicechat": "Proximity voice chat и его служебные команды.",
    "WorldEdit": "Инструменты редактирования мира, регионов и схем.",
    "WorldGuard": "Защита регионов, флагов и зон.",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, name: str, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Heading 1", 17, "0D5C46"),
        ("Heading 2", 14, "134E4A"),
        ("Heading 3", 12, "166534"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("CopiMine — справочник по плагинам и командам")
    set_run_font(run, "Calibri", 22, True, "0B3B2E")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(
        "Сначала идут самописные плагины CopiMine. После них — готовые серверные плагины. "
        "У каждого блока есть короткое описание и таблица всех доступных команд."
    )
    set_run_font(run, "Calibri", 10.5, False, "4B5563")


def parse_alias_value(value: str) -> list[str]:
    cleaned = value.strip().strip("'\"")
    if not cleaned:
        return []
    if cleaned.startswith("[") and cleaned.endswith("]"):
        cleaned = cleaned[1:-1]
        return [item.strip().strip("'\"") for item in cleaned.split(",") if item.strip()]
    return [cleaned]


def load_plugin_meta(text: str) -> dict:
    result: dict[str, object] = {"commands": {}, "permissions": {}}
    current_section: str | None = None
    current_command: str | None = None
    current_permission: str | None = None
    current_multiline_key: str | None = None

    for raw_line in text.splitlines():
        line = raw_line.rstrip("\n\r")
        if not line or line.lstrip().startswith("#"):
            continue
        indent = len(line) - len(line.lstrip(" "))
        stripped = line.strip()

        if indent == 0 and ":" in stripped:
            key, value = stripped.split(":", 1)
            key = key.strip()
            value = value.strip().strip("'\"")
            current_section = key
            current_command = None
            current_permission = None
            current_multiline_key = None
            if key in {"depend", "softdepend", "libraries"}:
                result[key] = []
            elif key not in {"commands", "permissions"}:
                result[key] = value
            continue

        if current_section == "commands":
            if indent == 2 and stripped.endswith(":"):
                current_command = stripped[:-1].strip()
                result["commands"][current_command] = {}
                current_multiline_key = None
                continue

            if indent >= 4 and current_command and ":" in stripped:
                key, value = stripped.split(":", 1)
                key = key.strip()
                value = value.strip()
                payload = result["commands"][current_command]
                current_multiline_key = None
                if key == "aliases":
                    aliases = parse_alias_value(value)
                    payload[key] = aliases
                    if not value:
                        current_multiline_key = key
                else:
                    payload[key] = value.strip("'\"")
                    if value in {"|", ">"}:
                        payload[key] = ""
                        current_multiline_key = key
                continue

            if indent >= 6 and current_command and stripped.startswith("-") and current_multiline_key == "aliases":
                payload = result["commands"][current_command]
                payload.setdefault("aliases", []).append(stripped.lstrip("-").strip().strip("'\""))
                continue

            if indent >= 6 and current_command and current_multiline_key in {"description", "usage"}:
                payload = result["commands"][current_command]
                existing = str(payload.get(current_multiline_key, ""))
                extra = stripped.strip("'\"")
                payload[current_multiline_key] = (existing + " " + extra).strip()
                continue

        if current_section == "permissions":
            if indent == 2 and stripped.endswith(":"):
                current_permission = stripped[:-1].strip()
                result["permissions"][current_permission] = {}
                continue
            if indent >= 4 and current_permission and ":" in stripped:
                key, value = stripped.split(":", 1)
                result["permissions"][current_permission][key.strip()] = value.strip().strip("'\"")
            continue

        if current_section in {"depend", "softdepend", "libraries"} and stripped.startswith("-"):
            result.setdefault(current_section, []).append(stripped.lstrip("-").strip())

    return result


def meta_from_plugin_yml(path: Path) -> dict:
    return load_plugin_meta(path.read_text(encoding="utf-8"))


def meta_from_plugin_jar(path: Path) -> dict | None:
    try:
        with zipfile.ZipFile(path) as jar:
            text = jar.read("plugin.yml").decode("utf-8")
    except Exception:
        return None
    return load_plugin_meta(text)


def build_custom_records() -> list[PluginRecord]:
    records: list[PluginRecord] = []
    for spec in CUSTOM_PLUGIN_SPECS:
        meta = meta_from_plugin_yml(spec["plugin_yml"])
        commands = list(CUSTOM_COMMANDS.get(spec["name"], []))
        if not commands:
            for command_name, payload in (meta.get("commands") or {}).items():
                commands.append(
                    CommandRow(
                        command=f"/{command_name}",
                        description=str(payload.get("description") or "Команда плагина."),
                        usage=str(payload.get("usage") or ""),
                        aliases=list(payload.get("aliases") or []),
                    )
                )
        records.append(
            PluginRecord(
                name=spec["name"],
                description=spec["description"],
                version=str(meta.get("version", "")),
                main_class=str(meta.get("main", "")),
                jar_name=f"{spec['name']}.jar",
                depends=list(meta.get("depend") or []),
                softdepends=list(meta.get("softdepend") or []),
                commands=commands,
            )
        )
    return records


def build_third_party_records() -> list[PluginRecord]:
    records: list[PluginRecord] = []
    for jar_path in sorted(SERVER_PLUGINS_DIR.glob("*.jar"), key=lambda item: item.name.lower()):
        if jar_path.name.startswith("CopiMine"):
            continue
        meta = meta_from_plugin_jar(jar_path)
        if not meta:
            continue
        name = str(meta.get("name") or jar_path.stem)
        commands: list[CommandRow] = []
        for command_name, payload in (meta.get("commands") or {}).items():
            description = str(payload.get("description") or "").strip()
            usage = str(payload.get("usage") or "").strip()
            aliases = list(payload.get("aliases") or [])
            if not description:
                description = f"Команда плагина {name}."
            commands.append(CommandRow(f"/{command_name}", description, usage, aliases))
        records.append(
            PluginRecord(
                name=name,
                description=THIRD_PARTY_DESCRIPTIONS.get(name, "Готовый серверный плагин, подключённый в сборку CopiMine."),
                version=str(meta.get("version", "")),
                main_class=str(meta.get("main", "")),
                jar_name=jar_path.name,
                depends=list(meta.get("depend") or []),
                softdepends=list(meta.get("softdepend") or []),
                commands=commands,
            )
        )
    return records


def add_plugin_meta_block(doc: Document, plugin: PluginRecord) -> None:
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(6)
    line = []
    if plugin.version:
        line.append(f"Версия: {plugin.version}")
    if plugin.jar_name:
        line.append(f"Файл: {plugin.jar_name}")
    if plugin.depends:
        line.append("Depend: " + ", ".join(plugin.depends))
    if plugin.softdepends:
        line.append("Softdepend: " + ", ".join(plugin.softdepends))
    meta_run = meta.add_run(" | ".join(line))
    set_run_font(meta_run, "Calibri", 9.5, False, "6B7280")


def configure_table(table) -> None:
    widths = [Inches(2.25), Inches(3.35), Inches(1.6), Inches(1.2)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, "Calibri", 9.5)


def add_command_table(doc: Document, commands: list[CommandRow]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Команда", "Что делает", "Как использовать", "Алиасы"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        set_cell_shading(cell, "E7F6EF")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, "Calibri", 10, True, "0B3B2E")
    for row in commands:
        cells = table.add_row().cells
        cells[0].text = row.command
        cells[1].text = row.description
        cells[2].text = row.usage or "—"
        cells[3].text = ", ".join(row.aliases) if row.aliases else "—"
    configure_table(table)


def add_plugin_section(doc: Document, plugin: PluginRecord) -> None:
    doc.add_heading(plugin.name, level=2)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(plugin.description)
    set_run_font(run, "Calibri", 10.5, False, "111827")
    add_plugin_meta_block(doc, plugin)

    if plugin.commands:
        add_command_table(doc, plugin.commands)
    else:
        no_cmd = doc.add_paragraph()
        no_cmd.paragraph_format.space_after = Pt(8)
        run = no_cmd.add_run("Отдельных чат-команд у этого плагина не заявлено. Он работает через GUI, блоки, API или как зависимость для других модулей.")
        set_run_font(run, "Calibri", 10, False, "374151")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def build_document(custom_records: Iterable[PluginRecord], third_party_records: Iterable[PluginRecord]) -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)

    section_intro = doc.add_paragraph()
    section_intro.paragraph_format.space_after = Pt(8)
    run = section_intro.add_run("Документ собран автоматически по plugin.yml, исходникам CopiMine и активным jar в папке сервера.")
    set_run_font(run, "Calibri", 9.5, False, "6B7280")

    doc.add_heading("Самописные плагины CopiMine", level=1)
    for plugin in custom_records:
        add_plugin_section(doc, plugin)

    doc.add_heading("Готовые плагины сервера", level=1)
    for plugin in third_party_records:
        add_plugin_section(doc, plugin)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


def main() -> None:
    custom_records = build_custom_records()
    third_party_records = build_third_party_records()
    out_path = build_document(custom_records, third_party_records)
    print(out_path)


if __name__ == "__main__":
    main()
